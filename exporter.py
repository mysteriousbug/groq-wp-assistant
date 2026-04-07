"""
UC-01 v2: Template-Driven DOCX Workpaper Exporter
Reads templates/workpaper_template.docx, fills by label matching.

Author: Ananya Aithal
"""

import os
import json
from docx import Document
from docx.shared import Pt
from io import BytesIO
from datetime import datetime
from db_models import ControlWorkpaper, TestingPhase, DocumentType


TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates", "workpaper_template.docx")


def _set_cell_value(cell, value: str):
    if not value:
        return
    ref_run = None
    for p in cell.paragraphs:
        for r in p.runs:
            ref_run = r
            break
        if ref_run:
            break

    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)

    first_para = cell.paragraphs[0]
    first_para.clear()

    lines = value.split("\n")
    for i, line in enumerate(lines):
        if i == 0:
            run = first_para.add_run(line)
        else:
            run = cell.add_paragraph().add_run(line)
        if ref_run:
            run.font.size = ref_run.font.size
            run.font.name = ref_run.font.name
            if ref_run.font.color and ref_run.font.color.rgb:
                run.font.color.rgb = ref_run.font.color.rgb
        else:
            run.font.size = Pt(8)


def _find_and_fill(table, label: str, value: str):
    label_clean = label.strip().lower().rstrip(":")
    for row in table.rows:
        cell_text = row.cells[0].text.strip().lower().rstrip(":")
        if cell_text == label_clean:
            _set_cell_value(row.cells[1], value)
            return True
    return False


def _build_cde_outcome(wp: ControlWorkpaper) -> str:
    """Build CDE Testing Outcome — walkthrough populates immediately."""
    wt_docs = [d for d in wp.documents
               if d.doc_type == DocumentType.TRANSCRIPT.value
               and d.phase == TestingPhase.WALKTHROUGH.value]
    cde_docs = [d for d in wp.documents
                if d.doc_type == DocumentType.TRANSCRIPT.value
                and d.phase == TestingPhase.CDE.value]
    support_docs = wp.get_supporting_docs()

    parts = ["CDE Source Data:"]

    if wt_docs:
        parts.append("Process Walkthrough:")
        parts.append("Discussion and walkthrough with the following stakeholders:")
        rcm = wp.rcm
        if rcm and rcm[0].get("control_ref"):
            owner = rcm[0].get("control_title", "")
            parts.append(f"- Control: {owner}")
        parts.append("")
        parts.append("Walkthrough transcripts reviewed:")
        for d in wt_docs:
            parts.append(f"- {d.filename} ({d.uploaded_at.strftime('%d %b %Y')})")
        parts.append("")
        for d in wt_docs:
            if d.summary:
                parts.append(f"Summary ({d.filename}):")
                parts.append(d.summary)
                parts.append("")

    if support_docs:
        parts.append("Supporting documents reviewed:")
        for d in support_docs:
            parts.append(f"- {d.filename} ({d.doc_type})")
        parts.append("")

    if cde_docs:
        parts.append("CDE discussion transcripts:")
        for d in cde_docs:
            parts.append(f"- {d.filename} ({d.uploaded_at.strftime('%d %b %Y')})")
        parts.append("")

    cde = wp.cde_result
    if cde:
        parts.append("CDE Assessment:")
        for s in cde.get("design_strengths", []):
            parts.append(f"- Strength: {s}")
        for g in cde.get("design_gaps", []):
            parts.append(f"- Gap: {g}")
        for c in cde.get("compensating_controls", []):
            parts.append(f"- Compensating: {c}")
        parts.append(f"Assessment: {cde.get('design_assessment', '')}")

    return "\n".join(parts) if len(parts) > 1 else ""


def export_workpaper(wp: ControlWorkpaper, template_path: str = None) -> BytesIO:
    path = template_path or TEMPLATE_PATH
    if not os.path.exists(path):
        raise FileNotFoundError(f"Template not found at {path}")

    doc = Document(path)
    if len(doc.tables) < 3:
        raise ValueError(f"Template has {len(doc.tables)} tables, expected 3+")

    rcm = wp.rcm
    rcm0 = rcm[0] if rcm else {}
    cde = wp.cde_result
    coe = wp.coe_result
    da = wp.da_result
    has_da = da is not None
    wt_docs = wp.get_transcripts(TestingPhase.WALKTHROUGH.value)
    has_wt = len(wt_docs) > 0

    # ── Table 0: Core Details ──
    core = {
        "Country": "Group-wide",
        "Legal Entity": "Standard Chartered Bank",
        "Risk Radar Themes": rcm0.get("risk_title", ""),
        "Assigned Team Member (CDE)": rcm0.get("audit_team_member", "Audit Team"),
        "Assigned Team Member (COE)": rcm0.get("audit_team_member", "Audit Team"),
        "Assigned Team Member (Substantive Testing)": "N/A",
        "Process": rcm0.get("process_title", ""),
        "Key Risk": rcm0.get("risk_description", ""),
        "Reference Number (Key Control)": rcm0.get("control_ref", ""),
        "Title (Key Control)": wp.control_name,
        "Due Date (Key Control)": datetime.now().strftime("%d %B %Y"),
        "Key Control Description": rcm0.get("control_description", ""),
        "CDE Required": "Yes" if (cde or has_wt) else "No",
        "Data Analytics (CDE)": "Yes" if has_da else "No",
        "Rationale for Skipping (CDE)": "N/A" if (cde or has_wt) else "Not yet performed",
        "COE Required": "Yes" if coe else "No",
        "Data Analytics (COE)": "Yes" if has_da else "No",
        "Rationale for Skipping (COE)": "N/A" if coe else "Not yet performed",
        "Substantive Test": "No",
        "Rationale for Skipping (Substantive Test)":
            "This is a controls testing and not a transactional testing, hence, substantive testing is not required.",
        "Data Analytics (Substantive Test)": "No",
    }
    for label, value in core.items():
        _find_and_fill(doc.tables[0], label, value)

    # ── Table 1: CDE ──
    cde_outcome = _build_cde_outcome(wp)
    cde_procedures = ""
    if has_wt or cde:
        cde_procedures = (
            "Reviewed control design documentation and conducted walkthrough with control owner "
            "to understand the design and operating procedures of the control. Assessed whether "
            "the control is appropriately designed to mitigate the identified risks."
        )

    proc_or_mon = "Monitoring"
    if rcm0.get("control_description", "") and "prevent" in rcm0.get("control_description", "").lower():
        proc_or_mon = "Processing"

    cde_conclusion = "Not Assessed"
    if cde:
        assess = cde.get("design_assessment", "")
        if "Well" in assess:
            cde_conclusion = "Effective"
        elif "Needs" in assess:
            cde_conclusion = "Partially Effective"
        elif "Poorly" in assess:
            cde_conclusion = "Ineffective"

    cde_map = {
        "Control Design Description": rcm0.get("control_description", ""),
        "Control Objective": rcm0.get("control_title", ""),
        "Control Frequency": "",
        "Applications Covered": "",
        "CDE Testing Procedures": cde_procedures,
        "CDE Testing Outcome (Results)": cde_outcome,
        "Processing or Monitoring Control?": proc_or_mon,
        "Manual or Automated Control?": "Manual",
        "Nature of Control": "Detective",
        "Complexity": "Moderate",
        "Testing completed by": "Audit Team",
        "Source Systems": "",
        "Other Data Sources / Systems": "Yes" if has_da else "No",
        "Other Data Sources / Systems Description":
            ", ".join(da.get("data_sources", [])) if has_da and da else "N/A",
        "Did analytics influence control outcome?": "Yes" if has_da else "No",
        "Analytics CT (control testing) coverage": "Full population" if has_da else "N/A",
        "Total Record Count": da.get("population_size", "N/A") if has_da and da else "N/A",
        "Analytics Procedure Description": da.get("analytics_performed", "N/A") if has_da and da else "N/A",
        "CDE Conclusion:": cde_conclusion,
        "Comments/Rationale": cde.get("conclusion", "") if cde else "",
    }
    for label, value in cde_map.items():
        _find_and_fill(doc.tables[1], label, value)

    # ── Table 2: COE ──
    coe_procedures, coe_outcome = "", ""
    if coe:
        coe_procedures = (
            "Select a representative sample within the audit period as per GIA methodology "
            f"and assess for appropriateness as per CDE.\n\nProcedure: {coe.get('testing_procedure', '')}"
        )
        parts = [f"COE Test Step\n{coe.get('results_summary', '')}"]
        for d in coe.get("deviation_details", []):
            parts.append(f"- {d}")
        coe_outcome = "\n".join(parts)

    sample_start, sample_end = "", ""
    if coe and coe.get("sample_period"):
        sp = coe["sample_period"]
        for sep in [" - ", " to ", " – "]:
            if sep in sp:
                sample_start, sample_end = sp.split(sep, 1)
                break
        else:
            sample_start = sp

    coe_conclusion = "Not Assessed"
    if wp.effectiveness != "Not Assessed":
        coe_conclusion = wp.effectiveness
    elif coe:
        devs = coe.get("deviations_found", 0)
        coe_conclusion = "Effective" if devs == 0 else ("Partially Effective" if devs <= 2 else "Ineffective")

    comments = coe.get("conclusion", "") if coe else ""
    exceptions = wp.exceptions
    if exceptions:
        comments += "\n\nExceptions identified:"
        for i, exc in enumerate(exceptions):
            comments += f"\n{i+1}. [{exc.get('severity', 'Medium')}] {exc.get('description', '')}"

    coe_map = {
        "COE Testing Procedures": coe_procedures,
        "COE Testing Outcome (Results)": coe_outcome,
        "Population Title": f"{wp.control_name} — Operating Evidence" if coe else "",
        "Population Description": f"Population of control execution records for {wp.control_name}." if coe else "",
        "Number of Items": coe.get("sample_size", "") if coe else "",
        "Source": "",
        "Relevance and Reliability": "Data from production systems verified against control documentation." if coe else "",
        "Stratification of Population": "Stratified by time period and platform." if coe else "",
        "Stratification Rationale": "Coverage across platforms and time periods." if coe else "",
        "Sample Size": coe.get("sample_size", "") if coe else "",
        "Sample Approach": "Representative sampling per GIA methodology" if coe else "",
        "Sample Start Date": sample_start,
        "Sample End Date": sample_end,
        "Complexity": "Moderate" if coe else "",
        "Testing completed by": "Audit Team" if coe else "",
        "Source Systems": "",
        "Other Data Sources / Systems": "Yes" if has_da else "No",
        "Other Data Sources / Systems Description":
            ", ".join(da.get("data_sources", [])) if has_da and da else "N/A",
        "Did analytics influence control outcome?": "Yes" if has_da else "No",
        "Analytics CT (control testing) coverage": "Full population" if has_da else "N/A",
        "Total Record Count": da.get("population_size", "N/A") if has_da and da else "N/A",
        "Analytics Procedure Description": da.get("analytics_performed", "N/A") if has_da and da else "N/A",
        "COE Conclusion:": coe_conclusion,
        "Comments/Rationale": comments,
    }
    for label, value in coe_map.items():
        _find_and_fill(doc.tables[2], label, value)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
